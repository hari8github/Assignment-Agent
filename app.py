from flask import Flask, render_template, request, jsonify, send_file
from main import create_enhanced_assignment
import os, json
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import textwrap

app = Flask(__name__, template_folder="templates", static_folder="static")

# Global variable to store the current assignment data
current_assignment = None

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/generate", methods=["POST"])
def generate_assignment():
    global current_assignment
    
    try:
        data = request.get_json()
        topic = data.get("topic")
        
        if not topic:
            return jsonify({"error": "No topic provided."}), 400

        result = create_enhanced_assignment(topic)
        output = result.get("output")
        
        if output is None:
            return jsonify({"error": "No output generated"}), 500
            
        try:
            current_assignment = json.loads(output)
            
            if not isinstance(current_assignment.get('sources'), list):
                current_assignment['sources'] = []
            
            return jsonify({"success": True, "data": current_assignment})
            
        except json.JSONDecodeError as e:
            return jsonify({"error": "Invalid JSON response from generator"}), 500

    except Exception as e:
        return jsonify({"error": str(e)}), 500

def format_content_as_text(assignment_data):
    """Convert assignment data to formatted text"""
    content = f"# {assignment_data['topic']}\n\n"
    content += f"Written by: {assignment_data['author']}\n"
    content += f"Date: {assignment_data['date']}\n\n"
    
    # Introduction
    content += f"## Introduction\n\n{assignment_data['introduction']}\n\n"
    
    # Main sections
    for i, section in enumerate(assignment_data['main_sections'], 1):
        content += f"## {i}. {section['title']}\n\n{section['content']}\n\n"
    
    # Conclusion
    content += f"## Conclusion\n\n{assignment_data['conclusion']}\n\n"
    
    # Sources
    if assignment_data.get('sources'):
        content += "## Sources\n\n"
        for i, source in enumerate(assignment_data['sources'], 1):
            content += f"{i}. {source}\n"
    
    return content

def create_txt_file(assignment_data, filepath):
    """Create TXT file from assignment data"""
    content = format_content_as_text(assignment_data)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

def create_pdf_file(assignment_data, filepath):
    """Create PDF file from assignment data"""
    doc = SimpleDocTemplate(filepath, pagesize=letter, topMargin=1*inch)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    story.append(Paragraph(assignment_data['topic'], title_style))
    story.append(Spacer(1, 20))
    
    # Author and date
    story.append(Paragraph(f"Written by: {assignment_data['author']}", styles['Normal']))
    story.append(Paragraph(f"Date: {assignment_data['date']}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Introduction
    story.append(Paragraph("Introduction", styles['Heading2']))
    story.append(Paragraph(assignment_data['introduction'], styles['Normal']))
    story.append(Spacer(1, 15))
    
    # Main sections
    for i, section in enumerate(assignment_data['main_sections'], 1):
        story.append(Paragraph(f"{i}. {section['title']}", styles['Heading2']))
        story.append(Paragraph(section['content'], styles['Normal']))
        story.append(Spacer(1, 15))
    
    # Conclusion
    story.append(Paragraph("Conclusion", styles['Heading2']))
    story.append(Paragraph(assignment_data['conclusion'], styles['Normal']))
    story.append(Spacer(1, 15))
    
    # Sources
    if assignment_data.get('sources'):
        story.append(Paragraph("Sources", styles['Heading2']))
        for i, source in enumerate(assignment_data['sources'], 1):
            story.append(Paragraph(f"{i}. {source}", styles['Normal']))
    
    doc.build(story)

def create_docx_file(assignment_data, filepath):
    """Create DOCX file from assignment data"""
    doc = Document()
    
    # Title
    title = doc.add_heading(assignment_data['topic'], 0)
    title.alignment = 1  # Center alignment
    
    # Author and date
    doc.add_paragraph(f"Written by: {assignment_data['author']}")
    doc.add_paragraph(f"Date: {assignment_data['date']}")
    doc.add_paragraph()  # Empty line
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(assignment_data['introduction'])
    
    # Main sections
    for i, section in enumerate(assignment_data['main_sections'], 1):
        doc.add_heading(f"{i}. {section['title']}", level=1)
        doc.add_paragraph(section['content'])
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(assignment_data['conclusion'])
    
    # Sources
    if assignment_data.get('sources'):
        doc.add_heading('Sources', level=1)
        for i, source in enumerate(assignment_data['sources'], 1):
            doc.add_paragraph(f"{i}. {source}")
    
    doc.save(filepath)

@app.route("/download/<format>")
def download_assignment(format):
    """Download assignment in original format (without edits)"""
    global current_assignment
    
    if not current_assignment:
        return jsonify({"error": "No assignment data available. Please generate an assignment first."}), 404

    filename = f"assignment.{format}"
    filepath = os.path.join(os.getcwd(), filename)

    try:
        if format == "txt":
            create_txt_file(current_assignment, filepath)
        elif format == "pdf":
            create_pdf_file(current_assignment, filepath)
        elif format == "docx":
            create_docx_file(current_assignment, filepath)
        else:
            return jsonify({"error": "Unsupported format"}), 400

        return send_file(filepath, as_attachment=True)
    
    except Exception as e:
        print(f"Error creating {format} file: {e}")
        return jsonify({"error": f"Failed to create {format} file"}), 500

@app.route("/download-edited", methods=["POST"])
def download_edited():
    """Download assignment with user edits"""
    try:
        data = request.get_json()
        format_type = data.get("format")
        
        # Create assignment data from edited content
        edited_assignment = {
            "topic": data.get("title", "Assignment"),
            "author": "AI Research Assistant",
            "date": data.get("date", ""),
            "introduction": data.get("introduction", ""),
            "main_sections": data.get("sections", []),
            "conclusion": data.get("conclusion", ""),
            "sources": current_assignment.get('sources', []) if current_assignment else []
        }

        filename = f"assignment_edited.{format_type}"
        filepath = os.path.join(os.getcwd(), filename)

        if format_type == "txt":
            create_txt_file(edited_assignment, filepath)
        elif format_type == "pdf":
            create_pdf_file(edited_assignment, filepath)
        elif format_type == "docx":
            create_docx_file(edited_assignment, filepath)
        else:
            return jsonify({"error": "Unsupported format"}), 400

        return send_file(filepath, as_attachment=True)
    
    except Exception as e:
        print(f"Error in download_edited: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/get-current-assignment")
def get_current_assignment():
    """API endpoint to get current assignment data"""
    global current_assignment
    if current_assignment:
        return jsonify(current_assignment)
    else:
        return jsonify({"error": "No assignment data available"}), 404

if __name__ == "__main__":
    app.run(debug=True)